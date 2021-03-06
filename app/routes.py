import json
from copy import deepcopy
from datetime import datetime
from flask import render_template, redirect, url_for, flash, request, abort, make_response, send_file
from app import app
from database.Models import *
from flask_login import login_user, logout_user, current_user, login_required
from sqlalchemy import update
from text import find_text
from urllib.parse import unquote
from werkzeug.utils import secure_filename
import os
import hashlib
import docx
import textract


@app.route('/')
@app.route('/index')
def index():
    return render_template('index.html')


@app.route('/stud/<int:id>', methods=['GET', 'POST'])
@login_required
def student_profile(id):
    if request.method == 'GET':
        user = Student.query.filter_by(id=id).first()
        if user is None:
            return redirect('https://stackoverflow.com/')

        if current_user.id == id:
            is_owner = True
        else:
            is_owner = False

        documents = Document.query.filter_by(student_id=id)

        notifications = Notification.query.filter_by(student_id=id).all()
        events = [Event.query.filter_by(id=item.event_id).first() for item in notifications]
        employers = [Employer.query.filter_by(id=event.employer_id).first() for event in
                     events]

        invites = [(item[0].name, item[0].id, item[1].name, item[1].id) for item in zip(events,
                                                                                        employers)]

        return render_template('student.html', id=id, name=user.name,
                               contacts=user.contacts,
                               cv=user.cv_hash, is_owner=is_owner, documents=documents,
                               invites=invites)
    else:
        pass


@app.route('/empl/<int:id>', methods=['GET', 'POST'])
@login_required
def employer_profile(id):
    if request.method == 'GET':
        user = Employer.query.filter_by(id=id).first()
        if user is None:
            return redirect('https://stackoverflow.com/')

        events = Event.query.filter_by(employer_id=id)
        events_tuples = [(event.id, event.name, event.description) for event in events]

        if current_user.id == id:
            is_owner = True
        else:
            is_owner = False

        return render_template('employer.html', name=user.name,
                               contacts=user.contacts,
                               description=user.description, events_tuples=events_tuples,
                               is_owner=is_owner)
    else:
        pass


@app.route('/stud/login', methods=['GET', 'POST'])
def stud_login():
    if request.method == "GET":
        res = make_response(render_template('login.html', page_title="Student Login",
                                            action="stud"))
        res.set_cookie("role", "student")
        return res
    else:
        email = request.form.get('Username')
        password = request.form.get('Password')

        user = Student.query.filter_by(login=email).first()
        if user is None or not user.check_password(password):
            return json.dumps({'error': 'Неправильный логин или пароль!'})

        login_user(user)
    return json.dumps({'success': user.id})


@app.route('/stud/signup', methods=['GET', 'POST'])
def stud_signup():
    if request.method == 'GET':
        res = make_response(render_template('signup.html', page_title="Student Sign Up",
                                            action='stud'))
        res.set_cookie("role", "student")
        return res
    else:
        username = request.form.get('Username')
        password = request.form.get('Password')
        email = request.form.get('Email')

        user = Student.query.filter_by(login=email).first()
        if user is not None:
            return json.dumps({'error': 'Пользователь с таким email уже существует.'})

        user = Student()
        user.name = username
        user.login = email
        user.set_password(password)
        db.session.add(user)
        db.session.commit()

        login_user(user)
        return json.dumps({'success': user.id})


@app.route('/empl/login', methods=['GET', 'POST'])
def empl_login():
    if request.method == "GET":
        res = make_response(render_template('login.html', page_title="Employer Login",
                                            action="empl"))
        res.set_cookie("role", "employer")
        return res
    else:
        username = request.form.get('Username')
        password = request.form.get('Password')

        user = Employer.query.filter_by(login=username).first()
        if user is None or not user.check_password(password):
            return json.dumps({'error': 'Неправильный логин или пароль!'})

        login_user(user)
    return json.dumps({'success': user.id})


@app.route('/empl/signup', methods=['GET', 'POST'])
def empl_signup():
    if request.method == 'GET':
        res = make_response(render_template('signup.html', page_title="Employer Sign Up",
                                            action='empl'))
        res.set_cookie("role", "employer")
        return res
    else:
        username = request.form.get('Username')
        password = request.form.get('Password')
        email = request.form.get('Email')

        user = Employer.query.filter_by(login=email).first()
        if user is not None:
            return json.dumps({'error': 'Пользователь с таким email уже существует.'})

        user = Employer()
        user.name = username
        user.login = email
        user.set_password(password)
        db.session.add(user)
        db.session.commit()

        login_user(user)
        return json.dumps({'success': user.id})


@app.route('/diploma/<string:name>', methods=['GET'])
def get_diploma(name):
    if os.path.splitext(name)[1] == '.pdf':
        return send_file(app.config['UPLOAD_FOLDER'] + name, mimetype='application/pdf')
    else:
        return send_file(app.config['UPLOAD_FOLDER'] + name, as_attachment=True)



@app.route('/upload_diploma/<int:id>', methods=['POST'])
def upload_diploma(id):
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        # if user does not select file, browser also
        # submit a empty part without filename
        if file.filename == '':
            return 'Файл не выбран', 400
        if os.path.splitext(file.filename)[1] != '.pdf' \
            and os.path.splitext(file.filename)[1] != '.docx':
            return 'Файлы данного расширения не поддерживаются', 400
        if file:
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], file.filename))
            filename = '{}{}'.format(hashlib.md5(file.read()).hexdigest(), os.path.splitext(file.filename)[1])
            os.rename('{}{}'.format(app.config['UPLOAD_FOLDER'], file.filename),
                      '{}{}'.format(app.config['UPLOAD_FOLDER'], filename))

            doc = Document()
            doc.student_id = id
            doc.filename = filename
            doc.link = '/diploma/' + filename
            doc.title = file.filename
            doc.type = 'Нет данных'
            doc.year = 2018
            doc.supervisor = 'Нет данных'
            text = textract.process('{}{}'.format(app.config['UPLOAD_FOLDER'], filename))

            if os.path.splitext(file.filename)[1] == '.pdf':
                text = textract.process('{}{}'.format(app.config['UPLOAD_FOLDER'], filename)).decode("utf-8")
            elif os.path.splitext(file.filename)[1] == '.docx':
                doc_file = docx.Document('{}{}'.format(app.config['UPLOAD_FOLDER'], filename))
                fullText = []
                for para in doc_file.paragraphs:
                    fullText.append(para.text)
                text = '\n'.join(fullText)

            doc.text = text

            db.session.add(doc)
            db.session.commit()

            data = {
                'id': doc.id,
                'speciality': 'Нет данных',
                'university': 'Нет данных',
                'supervisor': 'Нет данных',
                'type': 'Нет данных',
                'title': doc.title,
                'text': doc.text,
                'year': '2018',
                'link': doc.link,
            }
            find_text.put_diploma_to_index(data)

            return '', 200


@app.route('/remove_diploma/<int:id>', methods=['PATCH'])
def remove_diploma(id):
    doc = Document.query.filter_by(id=id).first()
    if doc.link.startswith('/diplomas/'):
        os.remove('{}{}'.format(app.config['UPLOAD_FOLDER'], doc.filename))
        find_text.remove_diploma_from_index(id)
    Document.query.filter_by(id=id).delete()
    db.session.commit()
    return '', 200


@app.route('/update_diploma/<int:id>', methods=['PATCH'])
def update_diploma(id):
    pass

@app.route('/empl/<int:id>', methods=['UPDATE'])
def update_empl(id):
    name = request.form.get('Username')
    contacts = request.form.get('Contacts')
    description = request.form.get('Description')

    emplr = Employer.query.filter_by(id=id).first()
    if emplr is None:
        return json.dumps({'error': 'Такого пользователя не существует.'})

    emplr.name = name
    emplr.contacts = contacts
    emplr.description = description

    db.session.commit()

    return json.dumps({'success': emplr.id})


@app.route('/stud/test')
def test_user_profile():
    return render_template('student.html',
                           name="Amazing Name",
                           contacts="t.me/absolutelyrandomguy",
                           cv="http://example.com",
                           diplomas=[
                               ("Тема диплома №1", "Описание диплома №1", "Данные диплома №1",
                                "Еще что-то диплома №1"),
                               ("Тема диплома №2", "Описание диплома №2", "Данные диплома №2",
                                "Еще что-то диплома №2"),
                               ("Тема диплома №3", "Описание диплома №3", "Данные диплома №3",
                                "Еще что-то диплома №3")
                           ])


@app.route('/success')
@login_required
def success():
    return render_template('test.html')


@app.route('/empl/theme/test')
def test_empl_theme():
    return render_template(
        'theme.html',
        project_students=[
            ("Name 1", 1),
            ("Name 2", 2),
        ],
        recommended_students=[
            ("Name Lastname", 3, 4.2),
            ("Another Nameless", 4, 4.9),
            ("Some RandomGuy", 5, 1)
        ]
    )


@app.route('/empl/<int:empl_id>/event/<int:id>', methods=['GET'])
def get_event(empl_id, id):
    event = Event.query.filter_by(id=id).first()
    employee = Employer.query.filter_by(id=empl_id).first()
    if event is None:
        return json.dumps({'error': 'Такого события не существует.'})

    event_students = [
        ("Name 1", 1),
        ("Name 2", 2),
    ]
    print(event.description.encode().decode('utf-8'))
    query = '{} {}'.format(str(event.name), str(event.description))
    print(query)
    res = find_text.find_students_by_theme(query)
    print(res)

    def enrich_response(record):
        doc = Document.query.filter_by(id=record['id']).first()
        if doc is None:
            return
        student = Student.query.filter_by(id=doc.student_id).first()

        event = Notification.query.filter_by(student_id=doc.student_id, event_id=id).first()

        if student is None:
            return {}

        return student.name, student.id, record['score'], False if event is None else True

    recommended_students = list(map(enrich_response, res))

    if current_user.id == empl_id:
        is_owner = True
    else:
        is_owner = False

    return render_template('theme.html', name=event.name, description=event.description,
                           company_name=employee.name, company_id=empl_id,
                           diploma=event.diploma,
                           event_students=event_students, recommended_students=recommended_students,
                           is_owner=is_owner)


@app.route('/empl/<int:empl_id>/event/<int:id>', methods=['UPDATE'])
def update_event(empl_id, id):
    event = Event.query.filter_by(id=id).first()
    if event is None:
        return json.dumps({'error': 'Такого события не существует.'})

    name = request.form.get('Name')
    description = request.form.get('Description')

    event = Event.query.filter_by(id=id).first()
    event.name = name
    event.description = description

    db.session.commit()

    return json.dumps({'success': id})


@app.route('/empl/<int:id>/event', methods=['POST'])
def post_event(id):
    name = request.form.get('Name')
    description = request.form.get('Description')
    is_diploma = bool(request.form.get('Diploma'))
    employer_id = id
    print('{} {} {} {}'.format(name, description, employer_id, is_diploma))
    event = Event()
    event.name = name
    event.description = description
    event.employer_id = employer_id
    event.diploma = is_diploma

    db.session.add(event)
    db.session.commit()

    find_text.add_event_to_index({"text": event.name + " " + event.description, "date": datetime.utcnow(),
                            "id": event.id})

    print('{} {} {} {}'.format(event.name, event.description, event.diploma, event.employer_id))
    return json.dumps({'success': event.id})


@app.route('/logout', methods=["POST"])
def logout():
    logout_user()
    return redirect(url_for("index"))


@app.route('/search', methods=['GET'])
def search():
    return render_template('search.html')


@app.route('/search/api', methods=['GET'])
def search_api__():
    query = request.args.get('q').lower()

    events = Event.query.all()

    filtered_events = filter(
        lambda event: query in event.name.lower() or query in event.description.lower(), events)

    def enrich_response(event):
        company_id = event.employer_id
        company_name = Employer.query.filter_by(id=company_id).first().name
        return {
            'id': event.id,
            'company_id': company_id,
            'company_name': company_name,
            'type': 1 if event.diploma else 2,
            'name': event.name,
            'description': event.description
        }

    return json.dumps(list(map(enrich_response, filtered_events)))


@app.route('/search_students_by_theme/api', methods=['GET'])
def search_api():
    query = request.args.get('q').lower()
    res = find_text.find_students_by_theme(query)

    def enrich_response(record):
        student_id = record['id']
        student = Student.query.filter_by(id=student_id).first()

        if student is None:
            return {}

        return student.name, student.id, record['score']

    return json.dumps(list(map(enrich_response, res)))


@app.route('/notification', methods=["POST"])
def send_notification():
    student_id = request.form.get('student_id')
    event_id = request.form.get('event_id')
    Notification.query.filter_by(student_id=student_id, event_id=event_id).delete()
    notification = Notification(student_id=student_id,
                                event_id=event_id,
                                checked=False)
    db.session.add(notification)
    db.session.commit()

    return json.dumps({'success': 'success'})


@app.route('/reject_notification', methods=['POST'])
def reject_notification():
    student_id = request.form.get('student_id')
    event_id = request.form.get('event_id')
    Notification.query.filter_by(student_id=student_id, event_id=event_id).delete()
    return json.dumps({'success': 'success'})


def import_vacancies():
    ids = {
        23: 'C#',
        22: 'финансы',
        21: 'еда',
        20: 'переезд',
        19: 'java python'
    }

    for key in ids:
        vacancies = find_text.find_event_by_string(ids[key])
        for v in vacancies:
            text = v['text']
            if len(text) > 1100:
                text = text[1:1100]

            event = Event()
            event.name = 'Стажировка'
            event.description = text
            event.employer_id = key
            event.diploma = False
            db.session.add(event)

    db.session.commit()
